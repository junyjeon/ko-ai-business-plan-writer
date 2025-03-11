from PyPDF2 import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
import io
import os
from docx import Document

def insert_text_to_pdf(input_pdf, output_pdf, text, position, font_size=12):
    """PDF 템플릿의 지정된 위치에 텍스트 삽입"""
    reader = PdfReader(input_pdf)
    writer = PdfWriter()
    
    # 텍스트 위치와 페이지 정보
    page_num, x, y = position
    
    # 해당 페이지 수정
    if page_num < len(reader.pages):
        # 텍스트 오버레이 생성
        packet = io.BytesIO()
        c = canvas.Canvas(packet, pagesize=A4)
        c.setFont("Helvetica", font_size)
        
        # 텍스트 줄바꿈 처리
        lines = text.split('\n')
        current_y = y
        line_height = font_size * 1.2
        
        for line in lines:
            c.drawString(x, current_y, line)
            current_y -= line_height
        
        c.save()
        packet.seek(0)
        
        # 오버레이 적용
        overlay = PdfReader(packet)
        page = reader.pages[page_num]
        page.merge_page(overlay.pages[0])
        
        # 수정된 페이지 추가
        writer.add_page(page)
    
    # 나머지 페이지 추가
    for i in range(len(reader.pages)):
        if i != page_num:
            writer.add_page(reader.pages[i])
    
    # 결과 저장
    with open(output_pdf, "wb") as f:
        writer.write(f)
    
    return output_pdf

def create_docx_with_section(section_content, output_path, title="1. 문제 인식 (Problem)_창업 아이템의 필요성"):
    """Word 문서에 섹션 내용 삽입"""
    doc = Document()
    
    # 제목 추가
    doc.add_heading(title, level=1)
    
    # 내용 분석
    lines = section_content.split('\n')
    for line in lines:
        if line.startswith('◦'):
            # 첫 번째 수준 항목
            p = doc.add_paragraph()
            p.add_run(line).bold = True
        elif line.startswith('-'):
            # 두 번째 수준 항목
            doc.add_paragraph(line, style='List Bullet')
        elif line.strip():
            # 일반 텍스트
            doc.add_paragraph(line)
    
    # 저장
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path

def create_docx_with_sections(output_path, sections_dict):
    """여러 섹션을 포함하는 Word 문서 생성"""
    doc = Document()
    
    # 각 섹션 추가
    for title, content in sections_dict.items():
        # 제목 추가
        doc.add_heading(title, level=1)
        
        # 내용 분석
        lines = content.split('\n')
        for line in lines:
            if line.startswith('◦'):
                # 첫 번째 수준 항목
                p = doc.add_paragraph()
                p.add_run(line).bold = True
            elif line.startswith('-'):
                # 두 번째 수준 항목
                doc.add_paragraph(line, style='List Bullet')
            elif line.strip():
                # 일반 텍스트
                doc.add_paragraph(line)
        
        # 섹션 간 간격 추가 (마지막 섹션이 아닌 경우)
        if title != list(sections_dict.keys())[-1]:
            doc.add_paragraph()
    
    # 저장
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path

def merge_docx_files(input_files, output_file):
    """여러 Word 문서를 하나로 병합"""
    # 첫 번째 문서를 기준으로 병합
    master_doc = Document(input_files[0])
    
    # 첫 번째 문서를 제외한 나머지 문서를 순회하며 내용 추가
    for i, file_path in enumerate(input_files[1:], 1):
        # 페이지 나누기 추가
        master_doc.add_page_break()
        
        # 다음 문서 열기
        doc = Document(file_path)
        
        # 모든 단락 추가
        for paragraph in doc.paragraphs:
            # 단락 스타일과 내용 복사
            p = master_doc.add_paragraph(paragraph.text)
            # 스타일 복사 (가능한 경우)
            try:
                p.style = paragraph.style
            except:
                pass
    
    # 저장
    master_doc.save(output_file)
    return output_file

def convert_docx_to_pdf(docx_path, pdf_path):
    """Word 문서를 PDF로 변환 (향후 구현)"""
    # 이 기능은 해당 라이브러리 설치 필요
    # 현재는 자동화가 힘들어 사용자에게 알림을 줌
    print(f"알림: Word 문서 '{docx_path}'를 PDF로 변환하려면 MS Word나 LibreOffice를 사용해주세요.")
    print(f"변환 후 '{pdf_path}'로 저장해주세요.")
    return pdf_path 